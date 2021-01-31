import redis    # 导入redis 模块

pool = redis.ConnectionPool(host='localhost', port=6379, decode_responses=True)
r = redis.Redis(host='localhost', port=6379, decode_responses=True, connection_pool=pool)
r.set('name', 'runoob', ex=3)  # 设置 name 对应的值
print(r.get('name'))  # 取出键 name 对应的值



a = """
                _._                                                  
           _.-``__ ''-._                                            
      _.-``    `.  `_.  ''-._           Redis 3.0.3 (255fcb1a/0) 64 bit
  .-`` .-```.  ```\/    _.,_ ''-._
 (    '      ,       .-`  | `,    )     Running in standalone mode
 |`-._`-...-` __...-.``-._|'` _.-'|     Port: 6379
 |    `-._   `._    /     _.-'    |     PID: 75858
  `-._    `-._  `-./  _.-'    _.-'                                  
 |`-._`-._    `-.__.-'    _.-'_.-'|                                  
 |    `-._`-._        _.-'_.-'    |           http://redis.io        
  `-._    `-._`-.__.-'_.-'    _.-'                                  
 |`-._`-._    `-.__.-'    _.-'_.-'|                                  
 |    `-._`-._        _.-'_.-'    |                                  
  `-._    `-._`-.__.-'_.-'    _.-'                                  
      `-._    `-.__.-'    _.-'                                      
          `-._        _.-'                                          
              `-.__.-'                                              
"""

print(a)

with open(r'E:\User\zhaoshuang\Desktop\a.docx', mode='rb') as f:
    a = f.read()
    b = a.decode(encoding='utf-8',errors='ignore')
    print(a)

#读取docx中的文本代码示例
import docx
#获取文档对象
file=docx.Document(r'E:\User\zhaoshuang\Desktop\y.docx')
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

#输出每一段的内容
for para in file.paragraphs:
    print(para.text)

#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)